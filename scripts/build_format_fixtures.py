#!/usr/bin/env python3
"""Build equivalent MD, TXT, DOCX, and PDF capability fixtures."""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Sequence


TITLE = "Skill Evolution Format Fixture"
METADATA = (
    ("Version", "1.2"),
    ("Effective date", "2026-07-26"),
    ("Replay budget", "13 runs"),
    ("Reference", "https://example.org/skill-evolution/policy"),
)
SECTIONS = (
    (
        "Decision",
        (
            "Automatic candidate replay must stop when sandbox preflight "
            "fails. It must never fall back to the host.",
        ),
    ),
    (
        "Constraints",
        (
            "Correctness and capability coverage are hard constraints.",
            "Every candidate and failed attempt remains visible.",
            "Only a human may approve release of a candidate skill.",
        ),
    ),
    (
        "Limitation",
        (
            "Visual preference is not inferred from deterministic metrics; "
            "screenshots require human review.",
        ),
    ),
)
TABLE_HEADERS = ("Dimension", "Baseline", "Guardrail")
TABLE_ROWS = (
    ("Correctness", "Required", "No regression"),
    ("Token usage", "Measured by component", "Report variance"),
    ("Duration", "Measured per run", "Report outliers"),
)


def _markdown() -> str:
    lines = [f"# {TITLE}", ""]
    for label, value in METADATA:
        lines.append(f"- **{label}:** {value}")
    for heading, paragraphs in SECTIONS[:2]:
        lines.extend(["", f"## {heading}", ""])
        lines.extend(paragraphs)
    lines.extend(
        [
            "",
            "## Metrics",
            "",
            "| " + " | ".join(TABLE_HEADERS) + " |",
            "| --- | --- | --- |",
        ]
    )
    lines.extend("| " + " | ".join(row) + " |" for row in TABLE_ROWS)
    heading, paragraphs = SECTIONS[2]
    lines.extend(["", f"## {heading}", "", *paragraphs, ""])
    return "\n".join(lines)


def _plain_text() -> str:
    lines = [TITLE, "=" * len(TITLE), ""]
    lines.extend(f"{label}: {value}" for label, value in METADATA)
    for heading, paragraphs in SECTIONS[:2]:
        lines.extend(["", heading.upper(), "-" * len(heading), *paragraphs])
    lines.extend(["", "METRICS", "-" * 7])
    lines.append(" | ".join(TABLE_HEADERS))
    lines.extend(" | ".join(row) for row in TABLE_ROWS)
    heading, paragraphs = SECTIONS[2]
    lines.extend(["", heading.upper(), "-" * len(heading), *paragraphs, ""])
    return "\n".join(lines)


def _build_docx(path: Path) -> None:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(10)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(TITLE)
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.bold = True

    for label, value in METADATA:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(value)

    for heading, paragraphs in SECTIONS[:2]:
        document.add_heading(heading, level=1)
        for text in paragraphs:
            document.add_paragraph(text)

    document.add_heading("Metrics", level=1)
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    widths = (2340, 3510, 3510)
    for index, text in enumerate(TABLE_HEADERS):
        cell = table.rows[0].cells[index]
        cell.text = text
        cell.width = Inches(widths[index] / 1440)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].bold = True
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F2F4F7")
        cell._tc.get_or_add_tcPr().append(shading)
    for row_values in TABLE_ROWS:
        cells = table.add_row().cells
        for index, text in enumerate(row_values):
            cells[index].text = text
            cells[index].width = Inches(widths[index] / 1440)
            cells[index].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

    table_properties = table._tbl.tblPr
    width_element = table_properties.first_child_found_in("w:tblW")
    width_element.set(qn("w:type"), "dxa")
    width_element.set(qn("w:w"), "9360")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    table_properties.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(widths[index]))

    heading, paragraphs = SECTIONS[2]
    document.add_heading(heading, level=1)
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def _build_pdf(path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    document = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=TITLE,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "FixtureTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.black,
        spaceAfter=10,
        alignment=0,
    )
    heading_style = ParagraphStyle(
        "FixtureHeading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=19,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=16,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "FixtureBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        leading=13,
        spaceAfter=6,
    )
    story = [Paragraph(TITLE, title_style)]
    for label, value in METADATA:
        story.append(
            Paragraph(f"<b>{label}:</b> {value}", body_style)
        )
    for heading, paragraphs in SECTIONS[:2]:
        story.append(Paragraph(heading, heading_style))
        story.extend(Paragraph(text, body_style) for text in paragraphs)
    story.append(Paragraph("Metrics", heading_style))
    table = Table(
        [TABLE_HEADERS, *TABLE_ROWS],
        colWidths=[1.625 * inch, 2.4375 * inch, 2.4375 * inch],
        repeatRows=1,
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEADING", (0, 0), (-1, -1), 11),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#A0A0A0")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend([table, Spacer(1, 4)])
    heading, paragraphs = SECTIONS[2]
    story.append(Paragraph(heading, heading_style))
    story.extend(Paragraph(text, body_style) for text in paragraphs)
    document.build(story)


def build(output_directory: Path) -> None:
    """Write all four file formats from one semantic data structure."""

    output_directory.mkdir(parents=True, exist_ok=True)
    (output_directory / "canonical.md").write_text(
        _markdown(),
        encoding="utf-8",
    )
    (output_directory / "canonical.txt").write_text(
        _plain_text(),
        encoding="utf-8",
    )
    _build_docx(output_directory / "canonical.docx")
    _build_pdf(output_directory / "canonical.pdf")


def _run_cli(arguments: Sequence[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--output",
        default="fixtures/document-formats",
    )
    options = parser.parse_args(arguments)
    output = Path(options.output).resolve()
    build(output)
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(_run_cli())
